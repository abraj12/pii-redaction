import docx
from typing import List, Dict, Tuple, Any
from fastapi import UploadFile
import io
import re
import fitz
import pytesseract
from PIL import Image

def _normalize_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\b(?:[A-Za-z]\s+){2,}[A-Za-z]\b", lambda m: m.group(0).replace(" ", ""), text)
    return text.strip()

def _iter_paragraphs(document: docx.Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in document.sections:
        for part in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph

def _replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    if not paragraph.runs:
        return
        
    logical_text_fast = "".join(run.text for run in paragraph.runs)
    
    # 1. Strip whitespace to handle normalization discrepancies between Node/Python
    raw_no_ws = []
    ws_map = []
    for i, char in enumerate(logical_text_fast):
        if not char.isspace():
            raw_no_ws.append(char)
            ws_map.append(i)
            
    raw_no_ws_str = "".join(raw_no_ws)
    
    matches = []
    for original, replacement in replacements.items():
        if not original:
            continue
            
        orig_no_ws = "".join(c for c in original if not c.isspace())
        if not orig_no_ws:
            continue
            
        search_start = 0
        while True:
            idx = raw_no_ws_str.find(orig_no_ws, search_start)
            if idx == -1:
                break
                
            raw_start = ws_map[idx]
            last_char_idx = idx + len(orig_no_ws) - 1
            raw_end = ws_map[last_char_idx] + 1
            
            matches.append({
                "start": raw_start,
                "end": raw_end,
                "original": original,
                "replacement": replacement
            })
            search_start = idx + len(orig_no_ws)
            
    if not matches:
        return
        
    # Sort matches by start ascending, then length descending
    matches.sort(key=lambda x: (x["start"], -(x["end"] - x["start"])))
    
    # Resolve overlaps
    valid_matches = []
    last_end = -1
    for m in matches:
        if m["start"] >= last_end:
            valid_matches.append(m)
            last_end = m["end"]
            
    # Sort valid matches by start DESCENDING (right-to-left)
    valid_matches.sort(key=lambda x: x["start"], reverse=True)
    
    # Build run mapping ONCE
    run_mapping = []
    for run_idx, run in enumerate(paragraph.runs):
        for char_idx, char in enumerate(run.text):
            run_mapping.append((run_idx, char_idx))
            
    # Apply replacements right-to-left
    for m in valid_matches:
        start_idx = m["start"]
        end_idx = m["end"] - 1
        replacement = m["replacement"]
        
        start_run_idx, start_char_idx = run_mapping[start_idx]
        end_run_idx, end_char_idx = run_mapping[end_idx]
        
        if start_run_idx == end_run_idx:
            run = paragraph.runs[start_run_idx]
            run.text = run.text[:start_char_idx] + replacement + run.text[end_char_idx + 1:]
        else:
            start_run = paragraph.runs[start_run_idx]
            start_run.text = start_run.text[:start_char_idx] + replacement
            
            for r_idx in range(start_run_idx + 1, end_run_idx):
                paragraph.runs[r_idx].text = ""
                
            end_run = paragraph.runs[end_run_idx]
            end_run.text = end_run.text[end_char_idx + 1:]

class DocumentProcessor:
    def __init__(self, pii_service):
        self.pii_service = pii_service
        self.MIN_PAGE_TEXT_CHARS = 50

    def _entities_from_results(self, text: str, analyzer_results: List[Any]) -> Dict[str, Any]:
        mapping = {}
        entities = []
        from app.services.pii_service import generate_fake

        for res in sorted(analyzer_results, key=lambda item: (item.start, -item.score)):
            entity_text = text[res.start:res.end]
            key = (res.entity_type, entity_text, res.start, res.end)
            if not entity_text or key in mapping:
                continue
            fake_val = generate_fake(res.entity_type, entity_text)
            mapping[key] = {"type": res.entity_type, "value": fake_val}
            entities.append({
                "type": res.entity_type,
                "text": entity_text,
                "fakeValue": fake_val,
                "score": float(res.score),
                "start": int(res.start),
                "end": int(res.end),
                "redact": True
            })

        breakdown = {}
        for ent in entities:
            breakdown[ent["type"]] = breakdown.get(ent["type"], 0) + 1

        return {"total_pii": len(entities), "breakdown": breakdown, "entities": entities}

    async def analyze_document_text(self, extracted_text: str, original_filename: str) -> Dict[str, Any]:
        text = _normalize_text(extracted_text)
        if not text:
            raise ValueError("PII_DETECTION_INPUT_EMPTY: extracted text was empty")
        analyzer_results = self.pii_service.analyze_text(text)
        stats = self._entities_from_results(text, analyzer_results)
        stats["extractedCharacters"] = len(text)
        return stats

    async def analyze_document(self, file: UploadFile, original_filename: str) -> Dict[str, Any]:
        ext = original_filename.split(".")[-1].lower()
        content = await file.read()
        extracted_text = ""
        analyzed_pages = 0

        if ext == "txt":
            extracted_text = _normalize_text(content.decode("utf-8"))
            analyzed_pages = 1
        elif ext == "pdf":
            document = fitz.open(stream=content, filetype="pdf")
            analyzed_pages = len(document)
            for page in document:
                page_text = page.get_text()
                if len(page_text.strip()) < self.MIN_PAGE_TEXT_CHARS:
                    # Perform OCR if native extraction yields very little text
                    try:
                        pix = page.get_pixmap(dpi=200)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        page_text = pytesseract.image_to_string(img)
                    except Exception as e:
                        # Fallback if tesseract isn't installed
                        print(f"OCR Failed or Unavailable: {e}")
                extracted_text += page_text + "\n"
        elif ext == "docx":
            document = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join(p.text for p in _iter_paragraphs(document) if p.text.strip())
            analyzed_pages = max(1, len(list(document.paragraphs)) // 10 + 1)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        stats = await self.analyze_document_text(extracted_text, original_filename)
        stats["analyzedPages"] = analyzed_pages
        return stats

    async def redact_document(self, file: UploadFile, original_filename: str, entities_to_redact: List[Dict]) -> Tuple[bytes, str]:
        ext = original_filename.split(".")[-1].lower()
        content = await file.read()
        replacements = {
            ent["text"]: ent["fakeValue"]
            for ent in entities_to_redact
            if ent.get("redact", True) and ent.get("text") and ent.get("fakeValue")
        }
        
        flattened_replacements = {}
        for original, fake_val in replacements.items():
            parts = [p.strip() for p in original.split('\n') if p.strip()]
            for part in parts:
                flattened_replacements[part] = fake_val
        replacements = flattened_replacements

        if ext == "txt":
            text = _normalize_text(content.decode("utf-8"))
            for original, replacement in replacements.items():
                text = text.replace(original, replacement)
            return text.encode("utf-8"), "txt"

        if ext == "pdf":
            document = fitz.open(stream=content, filetype="pdf")
            # Create a dictionary of {original_text: fake_text}
            for page in document:
                for original_text in replacements.keys():
                    # Native PDF redaction using PyMuPDF
                    # Find instances of the string
                    rects = page.search_for(original_text)
                    for rect in rects:
                        # Add a redaction annotation with a black fill
                        page.add_redact_annot(rect, fill=(0, 0, 0))
                # Apply the redactions, wiping out underlying pixels/text
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)
                
            out_stream = io.BytesIO()
            document.save(out_stream)
            return out_stream.getvalue(), "pdf"

        if ext == "docx":
            print(f"========== DOCX REDACTION FUNCTION ENTERED ==========")
            print(f"[REDACTION DEBUG]\nFUNCTION:\napp.extractors.document_processor.DocumentProcessor.redact_document")
            print(f"[REDACTION DEBUG]\nENTITY COUNT:\n{len(entities_to_redact)}")
            for ent in entities_to_redact:
                print(f"{ent.get('type')} start={ent.get('start')} end={ent.get('end')} length={len(ent.get('text', ''))}")
                
            document = docx.Document(io.BytesIO(content))
            for paragraph in _iter_paragraphs(document):
                _replace_in_paragraph(paragraph, replacements)
            out_stream = io.BytesIO()
            document.save(out_stream)
            out_bytes = out_stream.getvalue()
            
            import hashlib
            in_hash = hashlib.sha256(content).hexdigest()
            out_hash = hashlib.sha256(out_bytes).hexdigest()
            print(f"[REDACTION DEBUG]\nINPUT HASH:\n{in_hash}\nOUTPUT HASH:\n{out_hash}\nINPUT SIZE:\n{len(content)}\nOUTPUT SIZE:\n{len(out_bytes)}")
            
            # Step 6: Extract output text immediately
            test_markers = [
                "Rahul Sharma", "rahul.sharma@example.com", "+91 98765 43210", 
                "ABCDE1234F", "2345 6789 0123", "P1234567", "DL-0420110149646", 
                "123456789012", "HDFC0001234", "4111 1111 1111 1111", "192.168.1.100"
            ]
            reopened_doc = docx.Document(io.BytesIO(out_bytes))
            reopened_text = "\n".join(p.text for p in _iter_paragraphs(reopened_doc) if p.text.strip())
            for marker in test_markers:
                status = "PRESENT" if marker in reopened_text else "ABSENT"
                print(f"[REDACTION DEBUG]\nMARKER {marker}: {status}")

            return out_bytes, "docx"

        raise ValueError(f"Unsupported file type: {ext}")
